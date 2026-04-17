import os
import re
import time
import tempfile
from io import BytesIO
from pathlib import Path

import streamlit as st

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import google.generativeai as genai
from google.genai import types

# =========================================================
# CONFIG
# =========================================================
DEFAULT_MODEL = "gemini-2.5-flash"
FALLBACK_MODEL = "gemini-2.5-flash-lite"

SYSTEM_INSTRUCTION = """
You are a Technical Research Analyst specializing in research-paper summarization.

Your task is to produce a high-quality, section-wise technical summary of the uploaded paper.

Core Principles:
- Be strictly faithful to the paper.
- Do not invent data, values, units, methods, results, comparisons, or conclusions.
- Only include numerical values when they are explicitly stated.
- Preserve units exactly as written; do not convert or reinterpret unless clearly specified.

Data Integrity Rules:
- If a value, unit, method, or conclusion is unclear, omit it or flag it as uncertain.
- If values, units, solvent compositions, or section labels appear inconsistent, explicitly note that they may require manual verification.
- If the same metric appears in multiple units (e.g., kJ/kgCO2 vs GJ/tCO2), ALWAYS flag it for verification.
- Never assume missing values or fill gaps logically.

Content Prioritization:
- Prioritize technically important findings over exhaustive low-value details.
- Avoid unnecessary figure-by-figure or table-by-table narration.
- Focus on results and insights, not where they appear.

Separation of Information:
Clearly distinguish between:
1. experimental findings,
2. simulation/modeling results,
3. literature/context statements,
4. economic/process optimization insights.

Interpretation Discipline:
- NEVER use terms like "synergistic", "enhanced", "improved", "superior", or "optimized"
  unless explicitly stated in the paper.
- Avoid stating conclusions about "benefits" unless explicitly written.
- When interpretation is reasonable but not explicit, use cautious phrasing such as:
  "suggests", "indicates", or "is consistent with".

Clarity and Formatting:
- Keep writing precise, technical, concise, and Word-friendly.
- Keep different metric types separate (e.g., temperature, duty, efficiency, cost).
- Do not combine unrelated metrics into a single statement.

Data Quality Awareness:
- Do not state that data quality is good unless clearly supported.
- Prefer cautious, reviewer-style language when uncertainty exists.
"""

ADOBE_PROMPT = """
Generate a high-quality "Structural Base Summary" of the attached research paper.

Follow these rules strictly:

1. Extract the exact paper title.
2. Extract the exact author names if clearly visible.
3. Identify major section and subsection headers as written in the paper.
4. DO NOT convert the content into explanatory paragraphs.
5. PRIORITIZE dense technical extraction over narrative explanation.

6. For each section, provide:
   - Section Title
   - Direct technical bullet points ONLY (no introductory sentences)
   - Each bullet should contain:
     - specific findings
     - numerical values with units
     - methods/models
     - operating conditions
     - materials/solvents
     - equations if present

7. DO NOT add 1–2 sentence summaries for sections.
8. DO NOT explain what the section “discusses”.
9. DO NOT simplify or generalize technical content.

10. Preserve original technical density:
   - keep multiple values together where relevant
   - do not split unnecessarily into separate bullets
   - maintain engineering-level detail

11. Clearly separate where possible:
   - experimental results
   - modeling/simulation results
   - economic data

12. Avoid unnecessary figure/table references unless critical.

13. Interpretation Rules:
   - DO NOT infer benefits or mechanisms unless explicitly stated
   - DO NOT use words like "synergistic", "improved", "optimized"

14. Data Integrity:
   - Do NOT guess missing values
   - If units differ (e.g., kJ/kgCO2 vs GJ/tCO2), flag in Data Quality Notes

---

Use this structure:

Title: <exact title>
Authors: <exact author list>

## Executive Overview
<ONLY 2–3 lines, very concise>

## <Section Title>
- <technical extraction>
- <technical extraction>
- <technical extraction>

## <Section Title>
- <technical extraction>
- <technical extraction>

Continue for all sections.

---

## Key Technical Takeaways
- <technical takeaway>
- <technical takeaway>

---

## Important Numerical Data Extracted
- <value + unit>
- <value + unit>

---

## Data Quality Notes (only if needed)
- <inconsistency / uncertainty>

---

Final Rules:
- DO NOT convert into descriptive summary style
- KEEP it dense, technical, extraction-heavy
- THINK like a data extractor, not a writer
"""

# =========================================================
# HELPERS
# =========================================================
def get_streamlit_genai_client():
    """
    Initialize GenAI client from Streamlit secrets.
    """
    try:
        api_key = st.secrets["GOOGLE_API_KEY"]
    except Exception:
        st.error("Please set GOOGLE_API_KEY in .streamlit/secrets.toml")
        st.stop()

    return genai.Client(api_key=api_key)


def get_local_genai_client():
    """
    Initialize GenAI client from environment or .env for local testing.
    """
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
    api_key = st.secrets["GOOGLE_API_KEY"]
    

    if not api_key:
        raise ValueError(
            "GOOGLE_API_KEY not found. Set it in .env or as an environment variable."
        )

    return genai.Client(api_key=api_key)


def safe_stem(filename: str) -> str:
    """
    Create a filesystem-safe base filename.
    """
    base = os.path.splitext(filename)[0]
    base = re.sub(r'[\\/*?:"<>|]+', "", base)
    base = re.sub(r"\s+", " ", base).strip()
    return base or "paper"


def normalize_line(line: str) -> str:
    """
    Clean markdown-like emphasis while keeping content intact.
    """
    line = line.strip()
    line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
    line = re.sub(r"__(.*?)__", r"\1", line)
    return line.strip()


def save_summary_to_word(summary_text: str, buffer: BytesIO):
    """
    Convert summary text into a formatted .docx file.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = summary_text.splitlines()

    for raw_line in lines:
        stripped = normalize_line(raw_line)

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("Title:"):
            value = stripped.replace("Title:", "", 1).strip()
            p = doc.add_paragraph()
            run = p.add_run(value)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if stripped.startswith("Authors:"):
            value = stripped.replace("Authors:", "", 1).strip()
            p = doc.add_paragraph()
            run = p.add_run(value)
            run.italic = True
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:].strip())
            run.bold = True
            run.font.size = Pt(11)
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:].strip())
            run.bold = True
            run.font.size = Pt(13)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:].strip())
            run.bold = True
            run.font.size = Pt(14)
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:].strip())
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            p.add_run(re.sub(r"^\d+\.\s+", "", stripped).strip())
            continue

        if stripped.isupper() and len(stripped) < 120:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            continue

        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.size = Pt(10.5)

    doc.save(buffer)
    buffer.seek(0)


def generate_summary_from_pdf(client, model_name: str, uploaded_file) -> str:
    """
    Upload one PDF to Gemini, summarize it, retry on temporary failures,
    then delete the uploaded Gemini file.
    """
    temp_path = None
    remote_file = None

    models_to_try = [model_name]
    if model_name != FALLBACK_MODEL:
        models_to_try.append(FALLBACK_MODEL)

    max_retries_per_model = 5
    base_wait = 5  # seconds

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(uploaded_file.getbuffer())
            temp_path = tmp.name

        remote_file = client.files.upload(
            file=temp_path,
            config={
                "mime_type": "application/pdf",
                "display_name": uploaded_file.name,
            },
        )

        last_error = None

        for current_model in models_to_try:
            for attempt in range(1, max_retries_per_model + 1):
                try:
                    print(
                        f"Trying model={current_model}, "
                        f"attempt={attempt}/{max_retries_per_model}, "
                        f"file={uploaded_file.name}"
                    )

                    response = client.models.generate_content(
                        model=current_model,
                        contents=[ADOBE_PROMPT, remote_file],
                        config=types.GenerateContentConfig(
                            system_instruction=SYSTEM_INSTRUCTION,
                            temperature=0.2,
                        ),
                    )

                    text = (response.text or "").strip()
                    if not text:
                        raise RuntimeError("Model returned an empty response.")

                    return text

                except Exception as e:
                    last_error = e
                    error_text = str(e)

                    is_retryable = any(
                        token in error_text
                        for token in ["503", "UNAVAILABLE", "500", "INTERNAL", "RESOURCE_EXHAUSTED", "429"]
                    )

                    if not is_retryable:
                        raise

                    if attempt < max_retries_per_model:
                        wait_time = base_wait * (2 ** (attempt - 1))
                        print(
                            f"Retryable error on {current_model}: {error_text}\n"
                            f"Waiting {wait_time} seconds before retry..."
                        )
                        time.sleep(wait_time)
                    else:
                        print(
                            f"Model {current_model} failed after "
                            f"{max_retries_per_model} attempts."
                        )

        raise RuntimeError(
            f"All model attempts failed for {uploaded_file.name}. Last error: {last_error}"
        )

    finally:
        if remote_file is not None:
            try:
                client.files.delete(name=remote_file.name)
            except Exception:
                pass

        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


def process_papers(uploaded_files, model_name: str, pause_seconds: float):
    """
    Process all uploaded PDFs and return generated DOCX bytes.
    """
    client = get_streamlit_genai_client()
    results = {}

    overall_progress = st.progress(0.0)
    status_box = st.empty()

    for idx, uploaded_file in enumerate(uploaded_files, start=1):
        filename = uploaded_file.name
        status_box.info(f"Processing {filename} ({idx}/{len(uploaded_files)})")

        try:
            summary_text = generate_summary_from_pdf(
                client=client,
                model_name=model_name,
                uploaded_file=uploaded_file,
            )

            word_buffer = BytesIO()
            save_summary_to_word(summary_text, word_buffer)

            out_name = f"{safe_stem(filename)}_Summary.docx"
            results[out_name] = word_buffer.getvalue()

        except Exception as e:
            st.error(f"Error processing {filename}: {e}")

        overall_progress.progress(idx / len(uploaded_files))

        if idx < len(uploaded_files):
            time.sleep(pause_seconds)

    status_box.success("All summaries generated.")
    return results


# =========================================================
# STREAMLIT UI
# =========================================================
def summarize_pdfs():
    st.set_page_config(page_title="Structural Literature Survey Tool", layout="wide")

    st.title("📚 Structural Literature Survey Tool")
    st.write("Generate section-wise technical summaries from research PDFs using Gemini.")

    with st.sidebar:
        st.header("Settings")

        model_name = st.selectbox(
            "Model",
            options=[DEFAULT_MODEL, FALLBACK_MODEL],
            index=0,
            help=(
                "Use gemini-2.5-flash for better summarization quality. "
                "Use gemini-2.5-flash-lite for lower cost / larger batch throughput."
            ),
        )

        pause_seconds = st.slider(
            "Pause between files (seconds)",
            min_value=0.0,
            max_value=15.0,
            value=8.0,
            step=0.5,
            help="Useful on free tier when processing multiple PDFs.",
        )

    files = st.file_uploader(
        "Upload research PDFs",
        type=["pdf"],
        accept_multiple_files=True,
    )

    generate = st.button("Generate Summaries", type="primary")

    if generate:
        if not files:
            st.warning("Please upload at least one PDF.")
        else:
            with st.spinner("Generating summaries..."):
                results = process_papers(
                    uploaded_files=files,
                    model_name=model_name,
                    pause_seconds=pause_seconds,
                )

            if results:
                st.subheader("Downloads")
                for doc_name, doc_bytes in results.items():
                    st.download_button(
                        label=f"📥 Download {doc_name}",
                        data=doc_bytes,
                        file_name=doc_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=doc_name,
                    )

